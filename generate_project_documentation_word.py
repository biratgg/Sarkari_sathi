#!/usr/bin/env python3
"""
Generate comprehensive project documentation as Word document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import datetime
import os

def generate_word_documentation():
    """Generate comprehensive project documentation as Word document"""
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Sarkari Sathi', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.color.rgb = RGBColor(26, 35, 126)
    
    subtitle = doc.add_paragraph('RAG-based Government Information Chatbot')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor(92, 107, 192)
    subtitle_run.bold = True
    
    doc.add_paragraph(f'Project Documentation - Generated on: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', 1)
    toc_items = [
        "1. Project Overview",
        "2. Technologies and Tools",
        "3. API Keys and Configuration",
        "4. Project Structure",
        "5. Core Components",
        "6. Implementation Details",
        "7. Features and Capabilities",
        "8. Data Flow",
        "9. Deployment Information",
        "10. Technical Specifications",
        "11. Known Issues and Limitations"
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()
    
    # 1. Project Overview
    doc.add_heading('1. Project Overview', 1)
    doc.add_paragraph(
        'Sarkari Sathi is a Retrieval-Augmented Generation (RAG) chatbot designed to provide '
        'citizens with accurate, up-to-date information about government services. The system uses '
        'semantic vector search to retrieve relevant information from government documents and provides '
        'structured, source-grounded answers to user queries.'
    )
    doc.add_paragraph()
    
    doc.add_heading('Key Objectives:', 2)
    objectives = [
        "Enable 24/7 access to government service information",
        "Support multilingual queries (English and Nepali)",
        "Provide accurate, source-attributed answers",
        "Handle diverse document formats (PDF, DOCX, tables)",
        "Operate cost-effectively without external API dependencies"
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    doc.add_page_break()
    
    # 2. Technologies and Tools
    doc.add_heading('2. Technologies and Tools', 1)
    
    # Create table for technologies
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Technology/Tool'
    hdr_cells[1].text = 'Version'
    hdr_cells[2].text = 'Purpose'
    
    tech_data = [
        ['Python', '3.13', 'Core programming language'],
        ['FastAPI', '0.118.2', 'Web framework for REST API backend'],
        ['Uvicorn', '0.37.0', 'ASGI server for FastAPI'],
        ['Pinecone', '7.3.0', 'Vector database for semantic search'],
        ['Sentence Transformers', '5.1.1', 'Local embedding model (all-MiniLM-L6-v2)'],
        ['PyTorch', '2.8.0', 'Deep learning framework for embeddings'],
        ['pdfplumber', '0.11.8', 'PDF text extraction'],
        ['python-docx', '1.2.0', 'DOCX file processing'],
        ['python-dotenv', '1.1.1', 'Environment variable management'],
        ['Pydantic', '2.12.0', 'Data validation and settings'],
        ['Jinja2', '3.1.6', 'HTML templating'],
        ['NumPy', '2.3.3', 'Numerical computations'],
        ['Transformers', '4.57.0', 'Hugging Face transformers library'],
        ['Hugging Face Hub', '0.35.3', 'Model downloading and management']
    ]
    
    for tech, version, purpose in tech_data:
        row_cells = table.add_row().cells
        row_cells[0].text = tech
        row_cells[1].text = version
        row_cells[2].text = purpose
    
    doc.add_paragraph()
    doc.add_heading('Embedding Model:', 2)
    doc.add_paragraph('• Model: sentence-transformers/all-MiniLM-L6-v2', style='List Bullet')
    doc.add_paragraph('• Dimensions: 384', style='List Bullet')
    doc.add_paragraph('• Purpose: Converts text (queries and documents) into dense vector representations', style='List Bullet')
    doc.add_paragraph('• Advantages: Fast, lightweight, free to use, good semantic understanding', style='List Bullet')
    doc.add_page_break()
    
    # 3. API Keys and Configuration
    doc.add_heading('3. API Keys and Configuration', 1)
    
    doc.add_heading('Required API Keys:', 2)
    api_table = doc.add_table(rows=1, cols=4)
    api_table.style = 'Light Grid Accent 1'
    api_hdr = api_table.rows[0].cells
    api_hdr[0].text = 'Service'
    api_hdr[1].text = 'Environment Variable'
    api_hdr[2].text = 'Purpose'
    api_hdr[3].text = 'Where to Get'
    
    api_data = [
        ['Pinecone', 'PINECONE_API_KEY', 'Access to vector database', 'https://app.pinecone.io'],
        ['Pinecone', 'PINECONE_ENVIRONMENT', 'Pinecone environment/region', 'Pinecone dashboard'],
        ['OpenAI', 'OPENAI_API_KEY', 'Optional: for paid version', 'https://platform.openai.com']
    ]
    
    for service, env_var, purpose, where in api_data:
        row = api_table.add_row().cells
        row[0].text = service
        row[1].text = env_var
        row[2].text = purpose
        row[3].text = where
    
    doc.add_paragraph()
    doc.add_heading('Configuration File (config.py):', 2)
    config_items = [
        "PINECONE_API_KEY: Your Pinecone API key",
        "PINECONE_ENVIRONMENT: Your Pinecone environment",
        "PINECONE_INDEX_NAME: Name of the index (default: 'rag-chatbot')",
        "FREE_INDEX_NAME: Name for free version index (default: 'rag-chatbot-free')",
        "OPENAI_API_KEY: OpenAI API key (optional, for paid version)",
        "APP_HOST: Server host (default: '0.0.0.0')",
        "APP_PORT: Server port (default: 8000)",
        "EMBEDDING_MODEL: Embedding model name (default: 'text-embedding-ada-002')",
        "FREE_EMBEDDING_MODEL: Free model name (default: 'all-MiniLM-L6-v2')",
        "TOP_K_RESULTS: Number of documents to retrieve (default: 5)",
        "SIMILARITY_THRESHOLD: Minimum similarity score (default: 0.3)"
    ]
    for item in config_items:
        doc.add_paragraph(item, style='List Bullet')
    doc.add_page_break()
    
    # 4. Project Structure
    doc.add_heading('4. Project Structure', 1)
    
    structure_table = doc.add_table(rows=1, cols=2)
    structure_table.style = 'Light Grid Accent 1'
    struct_hdr = structure_table.rows[0].cells
    struct_hdr[0].text = 'File/Directory'
    struct_hdr[1].text = 'Purpose'
    
    structure_data = [
        ['main.py', 'FastAPI application entry point'],
        ['main_free.py', 'Free version (no OpenAI) entry point'],
        ['rag_chatbot_free.py', 'RAG chatbot logic with rule-based generation'],
        ['vector_store_free.py', 'Pinecone vector store operations'],
        ['config.py', 'Configuration and environment variables'],
        ['sample_data.py', 'Sample documents for testing'],
        ['setup_knowledge_base_free.py', 'Initialize knowledge base with sample data'],
        ['add_chandragiri_pdf.py', 'Add PDF documents to Pinecone'],
        ['add_nepal_docx.py', 'Add DOCX files to Pinecone'],
        ['templates/index.html', 'Web chat interface'],
        ['requirements.txt', 'Python dependencies'],
        ['.env', 'Environment variables (API keys)'],
        ['.venv/', 'Python virtual environment']
    ]
    
    for file_dir, purpose in structure_data:
        row = structure_table.add_row().cells
        row[0].text = file_dir
        row[1].text = purpose
    doc.add_page_break()
    
    # 5. Core Components
    doc.add_heading('5. Core Components', 1)
    
    doc.add_heading('5.1 Vector Store (vector_store_free.py)', 2)
    doc.add_paragraph('Purpose: Manages all interactions with Pinecone vector database')
    doc.add_paragraph()
    doc.add_paragraph('Key Methods:', style='List Bullet')
    doc.add_paragraph('• create_embedding(text): Converts text to 384-dimensional vector using MiniLM-L6-v2', style='List Bullet 2')
    doc.add_paragraph('• add_documents(documents): Embeds and stores documents in Pinecone', style='List Bullet 2')
    doc.add_paragraph('• search_similar(query, top_k): Finds most similar documents using cosine similarity', style='List Bullet 2')
    doc.add_paragraph('• get_index_stats(): Returns index statistics (vector count, dimensions)', style='List Bullet 2')
    doc.add_paragraph()
    doc.add_paragraph('Technical Details:', style='List Bullet')
    doc.add_paragraph('• Uses sentence-transformers for local embeddings (no API costs)', style='List Bullet 2')
    doc.add_paragraph('• Handles metadata size limits (40KB per vector)', style='List Bullet 2')
    doc.add_paragraph('• Supports chunking large documents automatically', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('5.2 RAG Chatbot (rag_chatbot_free.py)', 2)
    doc.add_paragraph('Purpose: Orchestrates the RAG pipeline: retrieval + generation')
    doc.add_paragraph()
    doc.add_paragraph('Key Methods:', style='List Bullet')
    doc.add_paragraph('• chat(query): Main entry point - processes user queries', style='List Bullet 2')
    doc.add_paragraph('• retrieve_relevant_documents(query): Searches Pinecone for relevant chunks', style='List Bullet 2')
    doc.add_paragraph('• generate_response(query, context): Rule-based answer extraction', style='List Bullet 2')
    doc.add_paragraph('• _handle_nepali_query(): Special handling for Nepali language queries', style='List Bullet 2')
    doc.add_paragraph()
    doc.add_paragraph('Response Extraction Types:', style='List Bullet')
    doc.add_paragraph('• Fee/Cost information', style='List Bullet 2')
    doc.add_paragraph('• Required documents (lists)', style='List Bullet 2')
    doc.add_paragraph('• Process/Steps', style='List Bullet 2')
    doc.add_paragraph('• Time/Duration', style='List Bullet 2')
    doc.add_paragraph('• Services available', style='List Bullet 2')
    doc.add_paragraph('• Location information', style='List Bullet 2')
    doc.add_paragraph('• Nepal-specific facts', style='List Bullet 2')
    doc.add_paragraph()
    
    doc.add_heading('5.3 FastAPI Backend (main_free.py)', 2)
    doc.add_paragraph('Purpose: RESTful API server for the chatbot')
    doc.add_paragraph()
    doc.add_paragraph('Endpoints:', style='List Bullet')
    doc.add_paragraph('• GET /: Serves the web chat interface', style='List Bullet 2')
    doc.add_paragraph('• POST /chat: Processes chat queries, returns responses', style='List Bullet 2')
    doc.add_paragraph('• POST /add-documents: Adds new documents to knowledge base', style='List Bullet 2')
    doc.add_paragraph('• GET /stats: Returns knowledge base statistics', style='List Bullet 2')
    doc.add_paragraph('• GET /health: Health check endpoint', style='List Bullet 2')
    doc.add_page_break()
    
    # 6. Implementation Details
    doc.add_heading('6. Implementation Details', 1)
    
    doc.add_heading('6.1 Document Processing Pipeline', 2)
    pipeline_steps = [
        "1. Ingestion: Documents are loaded from various sources (PDF, DOCX, text files)",
        "2. Text Extraction: pdfplumber extracts text from PDFs, python-docx from DOCX files",
        "3. Cleaning: Remove headers, footers, extra spaces, normalize formatting",
        "4. Chunking: Split large documents into smaller chunks (max 3000 chars) to avoid metadata limits",
        "5. Embedding: Each chunk is converted to 384-dimensional vector using MiniLM-L6-v2",
        "6. Indexing: Vectors stored in Pinecone with metadata (title, source, category, content preview)",
        "7. Metadata Limiting: Content in metadata limited to 2000 chars to stay under 40KB limit"
    ]
    for step in pipeline_steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_paragraph()
    
    doc.add_heading('6.2 Query Processing Flow', 2)
    query_steps = [
        "1. Query Reception: User query received via FastAPI /chat endpoint",
        "2. Preprocessing: Query is preprocessed (Nepali keyword expansion, normalization)",
        "3. Embedding: Query converted to vector using same embedding model",
        "4. Vector Search: Cosine similarity search in Pinecone (top-k=5, threshold=0.3)",
        "5. Context Assembly: Retrieved document chunks combined into context",
        "6. Query Classification: System identifies query type (fee, documents, process, etc.)",
        "7. Answer Extraction: Rule-based extractor pulls relevant information",
        "8. Response Formatting: Answer formatted with bullets, citations, source attribution",
        "9. Delivery: JSON response sent to frontend for display"
    ]
    for step in query_steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_page_break()
    
    # 7. Features and Capabilities
    doc.add_heading('7. Features and Capabilities', 1)
    
    doc.add_heading('7.1 Multilingual Support', 2)
    doc.add_paragraph('• Nepali Language Detection: Detects Devanagari script in queries', style='List Bullet')
    doc.add_paragraph('• Keyword Mapping: Maps Nepali question words to extraction methods', style='List Bullet')
    doc.add_paragraph('• Query Preprocessing: Adds English translations to Nepali queries for better embedding', style='List Bullet')
    doc.add_paragraph('• Supported Nepali Terms: कहाँ (where), कहिले (when), कति (how much), कसरी (how), के (what), कागजात (documents), शुल्क (fee), सेवा (service), etc.', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('7.2 Document Types Supported', 2)
    doc.add_paragraph('• PDF Files: Text extraction using pdfplumber, table extraction support', style='List Bullet')
    doc.add_paragraph('• DOCX Files: Full text extraction using python-docx', style='List Bullet')
    doc.add_paragraph('• Text Files: Direct text processing', style='List Bullet')
    doc.add_paragraph('• Tabular Data: Tables converted to structured text format', style='List Bullet')
    doc.add_paragraph('• Mixed Content: Handles documents with both text and tables', style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('7.3 Answer Extraction Capabilities', 2)
    extraction_capabilities = [
        "Fee Information: Extracts cost/fee details (including 'निःशुल्क' - free)",
        "Document Lists: Extracts numbered lists of required documents",
        "Process Steps: Identifies and extracts procedure information",
        "Time Information: Extracts duration, deadlines, service times",
        "Service Lists: Identifies available services from context",
        "Location Data: Extracts geographical and location information",
        "General Summaries: Fallback to general information extraction"
    ]
    for capability in extraction_capabilities:
        doc.add_paragraph(f"• {capability}", style='List Bullet')
    doc.add_page_break()
    
    # 8. Data Flow
    doc.add_heading('8. Data Flow', 1)
    
    doc.add_heading('8.1 Document Ingestion Flow', 2)
    doc.add_paragraph('Document → Text Extraction → Cleaning → Chunking → Embedding → Pinecone Storage')
    doc.add_paragraph()
    
    doc.add_heading('8.2 Query Processing Flow', 2)
    doc.add_paragraph('User Query → Preprocessing → Embedding → Pinecone Search → Context Retrieval → Answer Extraction → Response Formatting → User')
    doc.add_paragraph()
    
    doc.add_heading('8.3 Current Knowledge Base', 2)
    kb_table = doc.add_table(rows=1, cols=3)
    kb_table.style = 'Light Grid Accent 1'
    kb_hdr = kb_table.rows[0].cells
    kb_hdr[0].text = 'Document Type'
    kb_hdr[1].text = 'Count'
    kb_hdr[2].text = 'Source'
    
    kb_data = [
        ['Sample Documents', '10', 'sample_data.py'],
        ['Nepal Information', '1', 'nepal.docx'],
        ['Citizen Charter (Basic)', '1', 'Manual entry'],
        ['Citizen Charter (Complete PDF)', '151 chunks', 'नागरिक वडापत्र विषयक्षेत्रगतरुपमा.pdf'],
        ['Total Vectors', '163', 'Pinecone index: rag-chatbot-free']
    ]
    
    for doc_type, count, source in kb_data:
        row = kb_table.add_row().cells
        row[0].text = doc_type
        row[1].text = count
        row[2].text = source
    doc.add_page_break()
    
    # 9. Deployment Information
    doc.add_heading('9. Deployment Information', 1)
    
    doc.add_heading('9.1 Setup Steps', 2)
    setup_steps = [
        "Create virtual environment: python -m venv .venv",
        "Activate virtual environment: source .venv/bin/activate (Mac/Linux)",
        "Install dependencies: pip install -r requirements.txt",
        "Create .env file with API keys (PINECONE_API_KEY, PINECONE_ENVIRONMENT)",
        "Initialize knowledge base: python setup_knowledge_base_free.py",
        "Start server: python main_free.py",
        "Access web interface: http://localhost:8000"
    ]
    for step in setup_steps:
        doc.add_paragraph(step, style='List Number')
    doc.add_paragraph()
    
    doc.add_heading('9.2 Adding Documents', 2)
    add_doc_methods = [
        "PDF Files: python add_chandragiri_pdf.py [path_to_pdf]",
        "DOCX Files: python add_nepal_docx.py",
        "Via API: POST /add-documents with JSON payload",
        "Interactive: python interactive_add_documents.py"
    ]
    for method in add_doc_methods:
        doc.add_paragraph(f"• {method}", style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('9.3 System Requirements', 2)
    requirements = [
        "Python: 3.8 or higher",
        "RAM: Minimum 4GB (8GB recommended for embedding model)",
        "Storage: ~500MB for dependencies and models",
        "Internet: Required for Pinecone API and initial model download",
        "OS: macOS, Linux, or Windows"
    ]
    for req in requirements:
        doc.add_paragraph(f"• {req}", style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('9.4 Cost Analysis', 2)
    doc.add_paragraph('• Embedding Model: Free (local, no API calls)', style='List Bullet')
    doc.add_paragraph('• Pinecone: Free tier available (limited to 1 index, 100K vectors)', style='List Bullet')
    doc.add_paragraph('• Response Generation: Free (rule-based, no LLM API calls)', style='List Bullet')
    doc.add_paragraph('• Total Operating Cost: $0 (within free tier limits)', style='List Bullet')
    doc.add_paragraph('Note: For production scale, Pinecone paid plans start at $70/month', style='List Bullet')
    doc.add_page_break()
    
    # 10. Technical Specifications
    doc.add_heading('10. Technical Specifications', 1)
    
    doc.add_heading('10.1 Vector Database Configuration', 2)
    vector_specs = [
        "Index Name: rag-chatbot-free",
        "Dimensions: 384 (MiniLM-L6-v2)",
        "Metric: Cosine Similarity",
        "Cloud Provider: AWS",
        "Region: us-east-1",
        "Metadata Limit: 40KB per vector",
        "Current Vectors: 163"
    ]
    for spec in vector_specs:
        doc.add_paragraph(f"• {spec}", style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('10.2 API Endpoints', 2)
    endpoint_table = doc.add_table(rows=1, cols=4)
    endpoint_table.style = 'Light Grid Accent 1'
    ep_hdr = endpoint_table.rows[0].cells
    ep_hdr[0].text = 'Method'
    ep_hdr[1].text = 'Endpoint'
    ep_hdr[2].text = 'Purpose'
    ep_hdr[3].text = 'Request/Response'
    
    endpoint_data = [
        ['GET', '/', 'Web UI', 'Returns HTML chat interface'],
        ['POST', '/chat', 'Chat query', '{"message": "query"} → {"response": "...", "sources": [...]}'],
        ['POST', '/add-documents', 'Add docs', '{"documents": [...]} → {"success": true}'],
        ['GET', '/stats', 'Statistics', 'Returns index stats'],
        ['GET', '/health', 'Health check', 'Returns system status']
    ]
    
    for method, endpoint, purpose, req_resp in endpoint_data:
        row = endpoint_table.add_row().cells
        row[0].text = method
        row[1].text = endpoint
        row[2].text = purpose
        row[3].text = req_resp
    doc.add_page_break()
    
    # 11. Known Issues and Limitations
    doc.add_heading('11. Known Issues and Limitations', 1)
    
    doc.add_heading('11.1 Current Limitations', 2)
    limitations = [
        "Multilingual Support: Embedding model is English-centric, affecting Nepali query accuracy",
        "Response Quality: Rule-based extraction may miss nuanced queries",
        "Context Window: Limited to top-5 retrieved documents",
        "No Conversation Memory: Each query is independent, no context retention",
        "Metadata Size: Large documents must be chunked to fit 40KB metadata limit",
        "No OCR Support: Scanned PDFs with images not yet supported",
        "Similarity Threshold: Fixed at 0.3, may need tuning for different query types"
    ]
    for limitation in limitations:
        doc.add_paragraph(f"• {limitation}", style='List Bullet')
    doc.add_paragraph()
    
    doc.add_heading('11.2 Future Improvements', 2)
    improvements = [
        "Integrate multilingual embedding models (multilingual-MiniLM)",
        "Add OCR support for scanned documents",
        "Implement conversation memory and context retention",
        "Add hybrid search (keyword + semantic)",
        "Integrate local LLM for better response generation",
        "Add user authentication and query history",
        "Implement caching for frequently asked questions"
    ]
    for improvement in improvements:
        doc.add_paragraph(f"• {improvement}", style='List Bullet')
    
    # Save document
    doc_path = "Sarkari_Sathi_Project_Documentation.docx"
    doc.save(doc_path)
    print(f"✅ Word document generated successfully: {doc_path}")
    return doc_path

if __name__ == "__main__":
    generate_word_documentation()
